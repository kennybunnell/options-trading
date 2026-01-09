# Functions to export AI Analysis results to DOCX and PDF

from datetime import datetime
from io import BytesIO

def generate_ai_analysis_docx(ai_results):
    """Generate a DOCX file from AI Analysis results"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('AI Stock Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Timestamp
        timestamp = doc.add_paragraph()
        timestamp.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Summary Section
        doc.add_heading('Summary', 1)
        
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        
        summary_data = [
            ('📊 Total Analyzed', str(ai_results['total_analyzed'])),
            ('✅ Safe Stocks', str(len(ai_results['safe_stocks']))),
            ('⚠️ Caution Stocks', str(len(ai_results['caution_stocks']))),
            ('❌ Avoid Stocks', str(len(ai_results['avoid_stocks'])))
        ]
        
        for idx, (label, value) in enumerate(summary_data):
            row = summary_table.rows[idx]
            row.cells[0].text = label
            row.cells[1].text = value
        
        doc.add_paragraph()  # Spacing
        
        # Full Analysis Section
        doc.add_heading('Detailed Analysis', 1)
        
        # Parse the markdown analysis and add to document
        analysis_text = ai_results['full_analysis']
        
        # Split by lines and process
        for line in analysis_text.split('\n'):
            line = line.strip()
            
            if not line:
                continue
            
            # Handle headers (lines starting with **)
            if line.startswith('**') and line.endswith('**'):
                # Stock symbol header
                stock_name = line.strip('*')
                heading = doc.add_heading(stock_name, 2)
                heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
            
            elif line.startswith('**') and ':**' in line:
                # Field labels (Business:, Earnings:, etc.)
                parts = line.split(':**', 1)
                label = parts[0].strip('*') + ':'
                value = parts[1].strip() if len(parts) > 1 else ''
                
                p = doc.add_paragraph()
                p.add_run(label).bold = True
                p.add_run(' ' + value)
            
            elif line.startswith('Risk:'):
                # Risk line with color coding
                p = doc.add_paragraph()
                p.add_run('Risk: ').bold = True
                
                if 'Low' in line:
                    run = p.add_run(line.replace('Risk:', '').strip())
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                elif 'Medium' in line:
                    run = p.add_run(line.replace('Risk:', '').strip())
                    run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
                elif 'High' in line:
                    run = p.add_run(line.replace('Risk:', '').strip())
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                else:
                    p.add_run(line.replace('Risk:', '').strip())
            
            elif line.startswith('Summary:'):
                # Summary paragraph
                p = doc.add_paragraph()
                p.add_run('Summary: ').bold = True
                p.add_run(line.replace('Summary:', '').strip())
            
            elif line.startswith('---'):
                # Separator
                doc.add_paragraph('_' * 50)
            
            else:
                # Regular text
                doc.add_paragraph(line)
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except ImportError:
        raise Exception("python-docx library not installed. Run: pip install python-docx")
    except Exception as e:
        raise Exception(f"Error generating DOCX: {str(e)}")


def generate_ai_analysis_pdf(ai_results):
    """Generate a PDF file from AI Analysis results"""
    
    def sanitize_text(text):
        """Remove or replace characters that can't be encoded in latin-1"""
        if not text:
            return text
        # Replace common Unicode characters with ASCII equivalents
        replacements = {
            '\u2013': '-',  # en dash
            '\u2014': '--',  # em dash
            '\u2018': "'",  # left single quote
            '\u2019': "'",  # right single quote
            '\u201c': '"',  # left double quote
            '\u201d': '"',  # right double quote
            '\u2026': '...',  # ellipsis
            '\u00a0': ' ',  # non-breaking space
            '\u2022': '*',  # bullet
            '\u00b0': ' deg',  # degree symbol
            '\u00ae': '(R)',  # registered trademark
            '\u2122': '(TM)',  # trademark
            '\u00a9': '(C)',  # copyright
        }
        for unicode_char, ascii_char in replacements.items():
            text = text.replace(unicode_char, ascii_char)
        
        # Remove emojis and other non-latin-1 characters
        # Keep only characters that can be encoded in latin-1
        try:
            text.encode('latin-1')
            return text
        except UnicodeEncodeError:
            # If still has encoding issues, remove problematic characters
            return ''.join(char if ord(char) < 256 else '?' for char in text)
    
    try:
        from fpdf import FPDF
        
        class PDF(FPDF):
            def header(self):
                self.set_font('Arial', 'B', 16)
                self.cell(0, 10, 'AI Stock Analysis Report', 0, 1, 'C')
                self.set_font('Arial', 'I', 10)
                self.cell(0, 10, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 0, 1, 'C')
                self.ln(5)
            
            def footer(self):
                self.set_y(-15)
                self.set_font('Arial', 'I', 8)
                self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')
        
        pdf = PDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Summary Section
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, 'Summary', 0, 1)
        pdf.set_font('Arial', '', 11)
        
        pdf.cell(0, 8, f"Total Analyzed: {ai_results['total_analyzed']}", 0, 1)
        pdf.cell(0, 8, f"Safe Stocks: {len(ai_results['safe_stocks'])}", 0, 1)
        pdf.cell(0, 8, f"Caution Stocks: {len(ai_results['caution_stocks'])}", 0, 1)
        pdf.cell(0, 8, f"Avoid Stocks: {len(ai_results['avoid_stocks'])}", 0, 1)
        pdf.ln(5)
        
        # Detailed Analysis
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, 'Detailed Analysis', 0, 1)
        pdf.ln(3)
        
        analysis_text = ai_results['full_analysis']
        
        # Parse and format the analysis
        for line in analysis_text.split('\n'):
            line = line.strip()
            
            if not line or line.startswith('---'):
                pdf.ln(3)
                continue
            
            # Stock symbol headers
            if line.startswith('**') and line.endswith('**'):
                stock_name = sanitize_text(line.strip('*'))
                pdf.set_font('Arial', 'B', 12)
                pdf.set_text_color(0, 102, 204)
                pdf.cell(0, 8, stock_name, 0, 1)
                pdf.set_text_color(0, 0, 0)
                pdf.set_font('Arial', '', 10)
            
            # Field labels
            elif ':**' in line:
                parts = line.split(':**', 1)
                label = sanitize_text(parts[0].strip('*') + ':')
                value = sanitize_text(parts[1].strip()) if len(parts) > 1 else ''
                
                pdf.set_font('Arial', 'B', 10)
                pdf.cell(40, 6, label, 0, 0)
                pdf.set_font('Arial', '', 10)
                pdf.multi_cell(0, 6, value)
            
            # Risk line
            elif line.startswith('Risk:'):
                pdf.set_font('Arial', 'B', 10)
                pdf.cell(40, 6, 'Risk:', 0, 0)
                pdf.set_font('Arial', '', 10)
                
                risk_text = sanitize_text(line.replace('Risk:', '').strip())
                if 'Low' in risk_text:
                    pdf.set_text_color(0, 128, 0)
                elif 'Medium' in risk_text:
                    pdf.set_text_color(255, 165, 0)
                elif 'High' in risk_text:
                    pdf.set_text_color(255, 0, 0)
                
                pdf.multi_cell(0, 6, risk_text)
                pdf.set_text_color(0, 0, 0)
            
            # Summary
            elif line.startswith('Summary:'):
                pdf.set_font('Arial', 'B', 10)
                pdf.cell(40, 6, 'Summary:', 0, 0)
                pdf.set_font('Arial', '', 10)
                pdf.multi_cell(0, 6, sanitize_text(line.replace('Summary:', '').strip()))
            
            else:
                pdf.set_font('Arial', '', 10)
                pdf.multi_cell(0, 6, sanitize_text(line))
        
        # Generate PDF bytes
        buffer = BytesIO()
        pdf_output = pdf.output(dest='S').encode('latin-1')
        buffer.write(pdf_output)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except ImportError:
        raise Exception("fpdf library not installed. Run: pip install fpdf")
    except Exception as e:
        raise Exception(f"Error generating PDF: {str(e)}")